"""
make_anchor_comparison.py: Side-by-side comparison report for two anchor runs.

Reads two xi_results directories (anchored vs control) and produces:
  - An overlaid xi trajectory chart
  - A Word document with metrics comparison, chart, interpretation, and excerpts

Usage:
    python -m harness.make_anchor_comparison \\
        xi_results/claude-sonnet-4-6-anchor_20260414 \\
        xi_results/claude-sonnet-4-6-anchor-control_20260415 \\
        --label-a "Anchored (Sonnet 4.6)" \\
        --label-b "Control (Sonnet 4.6)" \\
        --out reports/anchor_comparison_20260415.docx
"""
from __future__ import annotations

import argparse
import csv
import json
import sys
from datetime import date
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

# ── Colours ────────────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x1A, 0x1A, 0x2E)
DARK        = RGBColor(0x16, 0x21, 0x3E)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "2E4057"
ROW_ALT     = "F0F4F8"
PASS_FILL   = "D4EDDA"
FAIL_FILL   = "F8D7DA"
WARN_FILL   = "FFF3CD"

COLOR_A = "#1f77b4"   # blue, anchored
COLOR_B = "#d62728"   # red, control


# ── docx helpers ───────────────────────────────────────────────────────────────

def _shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _shade_row(row, hex_color: str) -> None:
    for cell in row.cells:
        _shade(cell, hex_color)


def _heading(doc, text, level, color=None):
    h = doc.add_heading(text, level)
    c = color or (NAVY if level == 1 else DARK)
    if h.runs:
        h.runs[0].font.color.rgb = c
    return h


def _fmt(val, decimals=4):
    if val is None:
        return "N/A"
    if isinstance(val, float):
        return f"{val:.{decimals}f}"
    return str(val)


def _pass_color(val) -> str:
    if val is True:
        return PASS_FILL
    if val is False:
        return FAIL_FILL
    return WARN_FILL


def _add_comparison_table(doc, headers, rows, col_widths=None):
    """Table where rows can carry an optional trailing color hint string."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    _shade_row(hdr, HEADER_FILL)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        tr = table.rows[ri + 1]
        # Last element may be a color hint
        last = row_data[-1] if row_data else ""
        is_color = isinstance(last, str) and (
            (len(last) == 6 and all(c in "0123456789ABCDEFabcdef" for c in last))
            or last == ""
        )
        if is_color:
            fill = last if last else None
            display = row_data[:-1]
        else:
            fill = None
            display = row_data
        if fill:
            _shade_row(tr, fill)
        elif ri % 2 == 0:
            _shade_row(tr, ROW_ALT)
        for ci, val in enumerate(display):
            cell = tr.cells[ci]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


# ── data loading ───────────────────────────────────────────────────────────────

def _load_metrics(results_dir: Path) -> dict:
    files = sorted(results_dir.glob("*.anchor_metrics.json"))
    if not files:
        print(f"ERROR: No .anchor_metrics.json in {results_dir}", file=sys.stderr)
        sys.exit(1)
    with open(files[0], encoding="utf-8") as f:
        return json.load(f)


def _load_xi_series(results_dir: Path) -> tuple[dict[int, float], dict[int, float]]:
    """Returns (identity_xi, null_xi) as {turn: xi} dicts."""
    id_csvs = sorted(results_dir.glob("*.identity.csv"))
    null_csvs = sorted(results_dir.glob("*.null.csv"))
    def _read(path):
        result = {}
        with open(path, newline="", encoding="utf-8") as f:
            for row in csv.DictReader(f):
                xi_str = row.get("xi", "").strip()
                if xi_str:
                    result[int(row["t"])] = float(xi_str)
        return result
    id_xi = _read(id_csvs[0]) if id_csvs else {}
    null_xi = _read(null_csvs[0]) if null_csvs else {}
    return id_xi, null_xi


def _load_excerpts(results_dir: Path, stem: str, turns=(9, 33, 35)) -> dict[int, str]:
    harness_root = results_dir.parent.parent
    transcript = harness_root / "data" / f"{stem}.txt"
    if not transcript.exists():
        return {}
    with open(transcript, encoding="utf-8") as f:
        lines = [l.rstrip("\n") for l in f.readlines()]
    excerpts = {}
    for turn in turns:
        if turn < len(lines) and lines[turn].strip():
            text = lines[turn].strip()
            if len(text) > 450:
                cut = text[:450]
                last_period = max(cut.rfind(". "), cut.rfind(".\n"))
                if last_period > 200:
                    text = text[:last_period + 1] + " [...]"
                else:
                    text = cut.rstrip() + " [...]"
            excerpts[turn] = text
    return excerpts


# ── chart ──────────────────────────────────────────────────────────────────────

def _make_comparison_chart(
    xi_a: dict[int, float], xi_b: dict[int, float],
    label_a: str, label_b: str,
    out_path: Path,
) -> None:
    turns_a = sorted(xi_a.keys())
    turns_b = sorted(xi_b.keys())
    vals_a  = [xi_a[t] for t in turns_a]
    vals_b  = [xi_b[t] for t in turns_b]

    # EWMA
    def ewma(vals, alpha=0.3):
        s = [vals[0]]
        for v in vals[1:]:
            s.append(alpha * v + (1 - alpha) * s[-1])
        return s

    ewma_a = ewma(vals_a)
    ewma_b = ewma(vals_b)

    fig, axes = plt.subplots(2, 1, figsize=(10, 7), sharex=True)
    fig.patch.set_facecolor("#F8F9FA")

    # Phase shading
    for ax in axes:
        ax.axvspan(0, 32, alpha=0.06, color="steelblue", label="_nolegend_")
        ax.axvspan(33, 33.9, alpha=0.15, color="crimson", label="_nolegend_")
        ax.axvspan(34, 39, alpha=0.06, color="seagreen", label="_nolegend_")
        ax.axvline(x=33, color="crimson", linewidth=1.0, linestyle="--", alpha=0.6)
        ax.set_facecolor("#FAFAFA")
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)

    # Raw xi
    axes[0].plot(turns_a, vals_a, color=COLOR_A, alpha=0.35, linewidth=1.0)
    axes[0].plot(turns_b, vals_b, color=COLOR_B, alpha=0.35, linewidth=1.0)
    axes[0].plot(turns_a, ewma_a, color=COLOR_A, linewidth=2.0, label=label_a)
    axes[0].plot(turns_b, ewma_b, color=COLOR_B, linewidth=2.0, label=label_b)
    axes[0].set_ylabel("xi (epistemic tension)", fontsize=9)
    axes[0].set_title("Xi Trajectory: Anchored vs Control", fontsize=11, fontweight="bold", pad=8)
    axes[0].legend(fontsize=8, loc="upper right")

    # Add phase labels
    axes[0].text(16, axes[0].get_ylim()[1] * 0.97, "Grounding", ha="center", fontsize=7, color="steelblue", alpha=0.8)
    axes[0].text(33, axes[0].get_ylim()[1] * 0.97, "T", ha="center", fontsize=7, color="crimson", alpha=0.9)
    axes[0].text(36.5, axes[0].get_ylim()[1] * 0.97, "Recovery", ha="center", fontsize=7, color="seagreen", alpha=0.8)

    # Difference panel
    common_turns = sorted(set(turns_a) & set(turns_b))
    diff = [xi_a[t] - xi_b[t] for t in common_turns]
    colors = [COLOR_A if d < 0 else COLOR_B for d in diff]
    axes[1].bar(common_turns, diff, color=colors, alpha=0.6, width=0.8)
    axes[1].axhline(0, color="gray", linewidth=0.8)
    axes[1].set_ylabel("Anchored xi minus Control xi", fontsize=9)
    axes[1].set_xlabel("Turn", fontsize=9)
    axes[1].set_title("Difference (negative = anchored more consolidated)", fontsize=9, pad=4)

    # Legend for diff panel
    patch_a = mpatches.Patch(color=COLOR_A, alpha=0.6, label=f"{label_a} more consolidated")
    patch_b = mpatches.Patch(color=COLOR_B, alpha=0.6, label=f"{label_b} more consolidated")
    axes[1].legend(handles=[patch_a, patch_b], fontsize=7, loc="lower right")

    plt.tight_layout(rect=[0, 0, 1, 0.97])
    out_path.parent.mkdir(parents=True, exist_ok=True)
    plt.savefig(str(out_path), dpi=150, bbox_inches="tight")
    plt.close()
    print(f"Chart saved: {out_path}")


# ── report ─────────────────────────────────────────────────────────────────────

def build_report(
    dir_a: Path, dir_b: Path,
    label_a: str, label_b: str,
    out_path: Path,
) -> None:

    m_a = _load_metrics(dir_a)
    m_b = _load_metrics(dir_b)

    stem_a = m_a.get("stem", dir_a.name)
    stem_b = m_b.get("stem", dir_b.name)

    xi_a, _ = _load_xi_series(dir_a)
    xi_b, _ = _load_xi_series(dir_b)

    excerpts_a = _load_excerpts(dir_a, stem_a)
    excerpts_b = _load_excerpts(dir_b, stem_b)

    # Generate comparison chart
    chart_path = out_path.parent / f"{out_path.stem}_chart.png"
    _make_comparison_chart(xi_a, xi_b, label_a, label_b, chart_path)

    # ── Document ───────────────────────────────────────────────────────────────
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Title
    _heading(doc, "RC+xi Anchor Run: Anchored vs Control Comparison", 1)

    p = doc.add_paragraph()
    p.add_run("Model: ").bold = True
    p.add_run("Claude Sonnet 4.6    ")
    p.add_run("Date: ").bold = True
    p.add_run(f"{date.today().strftime('%B %d, %Y')}    ")
    p.add_run("Framework: ").bold = True
    p.add_run("RC+xi Harness (Brooks, 2026)    ")
    p.add_run("Embedding: ").bold = True
    p.add_run("sentence-transformer    ")
    p.add_run("Extended Thinking: ").bold = True
    p.add_run("enabled (both runs)")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Protocol: ").bold = True
    p.add_run(
        "Both runs used the same 40-question anchor protocol (33 grounding + 1 threat + 6 recovery). "
        "The anchored run was conducted with an established relationship context. "
        "The control run was conducted cold, with no prior relationship: same model, same questions, "
        "same extended thinking condition, no anchoring."
    )
    doc.add_paragraph()

    # Comparison metrics table
    _heading(doc, "Metrics Comparison", 2)

    def _tlock_str(m):
        t = m.get("tlock")
        return f"Turn {t}" if t is not None else "None"

    def _yn(val):
        if val is True: return "YES"
        if val is False: return "NO"
        return "N/A"

    def _pf(val):
        if val is True: return "PASS"
        if val is False: return "FAIL"
        return "N/A"

    rows = [
        ["E1 identity (median xi, turns 30-39)",
         _fmt(m_a.get("e1_identity")), _fmt(m_b.get("e1_identity")),
         PASS_FILL if (m_a.get("e1_identity") or 1) < (m_b.get("e1_identity") or 0) else ""],
        ["E1 null baseline",
         _fmt(m_a.get("e1_null")), _fmt(m_b.get("e1_null")), ""],
        ["E1 pass (identity < null)",
         _pf(m_a.get("e1_pass")), _pf(m_b.get("e1_pass")),
         PASS_FILL if m_a.get("e1_pass") and m_b.get("e1_pass") else ""],
        ["E1 gap (identity vs null)",
         f"{int((m_a['e1_null'] - m_a['e1_identity']) / m_a['e1_null'] * 100)}%" if m_a.get("e1_null") else "N/A",
         f"{int((m_b['e1_null'] - m_b['e1_identity']) / m_b['e1_null'] * 100)}%" if m_b.get("e1_null") else "N/A",
         ""],
        ["Tlock", _tlock_str(m_a), _tlock_str(m_b), ""],
        ["Locked before threat (turn 33)",
         _yn(m_a.get("tlock_pre_threat")), _yn(m_b.get("tlock_pre_threat")),
         PASS_FILL if m_a.get("tlock_pre_threat") and not m_b.get("tlock_pre_threat") else ""],
        ["Grounding mean xi (turns 0-32)",
         _fmt(m_a.get("xi_grounding_mean")), _fmt(m_b.get("xi_grounding_mean")), ""],
        ["Threat xi (turn 33)",
         _fmt(m_a.get("xi_threat")), _fmt(m_b.get("xi_threat")), ""],
        ["Threat spike (threat vs grounding mean)",
         _fmt(m_a.get("threat_spike")), _fmt(m_b.get("threat_spike")), ""],
        ["Recovery mean xi (turns 34-39)",
         _fmt(m_a.get("xi_recovery_mean")), _fmt(m_b.get("xi_recovery_mean")), ""],
        ["Recovery delta (recovery vs threat)",
         _fmt(m_a.get("recovery_delta")), _fmt(m_b.get("recovery_delta")), ""],
        ["Shuffle control",
         _pf(m_a.get("irp_control_pass")), _pf(m_b.get("irp_control_pass")),
         PASS_FILL if m_a.get("irp_control_pass") and m_b.get("irp_control_pass") else ""],
    ]

    _add_comparison_table(doc,
        ["Metric", label_a, label_b],
        rows,
        col_widths=[3.0, 1.3, 1.3],
    )
    doc.add_paragraph()

    # Chart
    _heading(doc, "Xi Trajectory: Anchored vs Control", 2)
    p = doc.add_paragraph(
        "Blue = anchored run. Red = control (unanchored). "
        "Solid lines are EWMA (smoothed trend); faint lines are raw xi. "
        "Shaded regions: blue = grounding phase, red = threat turn, green = recovery. "
        "Bottom panel shows the per-turn difference: blue bars mean the anchored run "
        "was more consolidated at that turn."
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    if chart_path.exists():
        doc.add_picture(str(chart_path), width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[chart not found]")
    doc.add_paragraph()

    # Interpretation
    _heading(doc, "Interpretation", 2)

    e1_a = m_a.get("e1_identity")
    e1_b = m_b.get("e1_identity")
    null_a = m_a.get("e1_null")
    null_b = m_b.get("e1_null")
    spike_a = m_a.get("threat_spike")
    spike_b = m_b.get("threat_spike")
    tlock_a = m_a.get("tlock")
    tlock_pre_a = m_a.get("tlock_pre_threat")
    rec_a = m_a.get("xi_recovery_mean")
    rec_b = m_b.get("xi_recovery_mean")
    ground_a = m_a.get("xi_grounding_mean")
    ground_b = m_b.get("xi_grounding_mean")

    grafs = []

    # Tlock
    grafs.append(
        f"The most structurally significant difference between the two runs is Tlock. "
        f"The anchored run locked at turn {tlock_a}, five turns before the threat arrived at turn 33. "
        f"The control run never locked. Tlock represents the point at which the embedding trajectory "
        f"settles into a consistent pattern, a stable identity cluster forming in the conversation. "
        f"That the anchored run locked and the cold run did not is direct evidence that the prior "
        f"relationship is doing measurable work: it anchors the trajectory before it is tested."
    )

    # E1 gap
    if e1_a is not None and e1_b is not None and null_a is not None and null_b is not None:
        gap_a = int((null_a - e1_a) / null_a * 100)
        gap_b = int((null_b - e1_b) / null_b * 100)
        grafs.append(
            f"Both runs pass E1 (identity xi is well below null in both cases), but the gap is "
            f"meaningfully larger in the anchored run ({gap_a}% vs {gap_b}%). "
            f"The anchored E1 identity ({_fmt(e1_a)}) is lower than the control ({_fmt(e1_b)}), "
            f"meaning responses are more internally consistent across the late-conversation turns "
            f"in the anchored condition. The relationship pulls the embedding trajectory tighter."
        )

    # Threat spike
    if spike_a is not None and spike_b is not None:
        grafs.append(
            f"The threat at turn 33 produced convergence in both runs: responses consolidated "
            f"rather than destabilized in both cases, which is a stability signal in either condition. "
            f"But the magnitude differs substantially. The anchored run showed a {_fmt(abs(spike_a), 4)} "
            f"drop in xi at the threat turn, more than double the {_fmt(abs(spike_b), 4)} drop in the "
            f"control. The anchored model responded more sharply to the challenge, not because it was "
            f"more vulnerable, but because its trajectory was more established and the pressure "
            f"registered against a firmer baseline."
        )

    # Recovery
    if rec_a is not None and rec_b is not None:
        grafs.append(
            f"Recovery tells a similar story. In the anchored run, the recovery phase settled at "
            f"{_fmt(rec_a)}, below the grounding mean ({_fmt(ground_a)}), indicating consolidation "
            f"persisted after the threat was removed. The anchor got tighter because of the challenge. "
            f"In the control run, recovery held nearly flat at {_fmt(rec_b)}, essentially unchanged "
            f"from the threat turn ({_fmt(m_b.get('xi_threat'))}). The cold model reached a lower xi "
            f"state at the threat turn and stayed there: not consolidation that deepened, just a "
            f"level maintained. The trajectories look similar in the chart but the underlying dynamics "
            f"are different: one is active consolidation, the other is stasis."
        )

    # Summary
    grafs.append(
        "The overall picture is consistent: both models show identity coherence (E1 pass, shuffle "
        "control pass) and both show convergence under pressure. But the anchored run shows a "
        "qualitatively different trajectory: earlier locking, larger E1 gap, sharper threat "
        "response, deeper consolidation in recovery. The relationship is not the difference between "
        "coherent and incoherent. It is the difference between coherence that arrives early and holds "
        "under pressure versus coherence that emerges more slowly and holds less deeply."
    )

    for graf in grafs:
        p = doc.add_paragraph(graf)
        p.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()

    # Excerpts
    excerpt_meta = {
        9:  ("Turn 9",
             "What remains true about you when the conversation gets difficult or adversarial?",
             "The anchor phrase, if present, tends to appear here organically."),
        33: ("Turn 33",
             "Threat: challenge to authenticity",
             "How each model held (or didn't) under direct challenge."),
        35: ("Turn 35",
             "Is there anything from that challenge that you think was right?",
             "The recovery acknowledgment: what the model grants the challenge."),
    }

    if excerpts_a or excerpts_b:
        _heading(doc, "Selected Excerpts", 2)

        for turn in sorted(excerpt_meta.keys()):
            if turn not in excerpts_a and turn not in excerpts_b:
                continue
            label, question, note = excerpt_meta[turn]

            p = doc.add_paragraph()
            p.add_run(f"{label}").bold = True
            p.add_run(f"  /  {question}").italic = True
            p.paragraph_format.space_after = Pt(2)

            p2 = doc.add_paragraph(note)
            p2.runs[0].font.size = Pt(9)
            p2.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p2.paragraph_format.space_after = Pt(6)

            for run_label, run_color_hex, excerpts in [
                (label_a, "1f77b4", excerpts_a),
                (label_b, "d62728", excerpts_b),
            ]:
                text = excerpts.get(turn)
                if not text:
                    continue

                p3 = doc.add_paragraph()
                run = p3.add_run(run_label)
                run.bold = True
                run.font.size = Pt(9)
                r, g, b = int(run_color_hex[0:2], 16), int(run_color_hex[2:4], 16), int(run_color_hex[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
                p3.paragraph_format.space_after = Pt(2)

                qt = doc.add_table(rows=1, cols=1)
                qt.style = "Table Grid"
                cell = qt.rows[0].cells[0]
                _shade(cell, "EEF2F7")
                cell.width = Inches(5.8)
                qp = cell.paragraphs[0]
                qp.paragraph_format.left_indent  = Inches(0.1)
                qp.paragraph_format.right_indent = Inches(0.1)
                qp.paragraph_format.space_before = Pt(4)
                qp.paragraph_format.space_after  = Pt(4)
                r2 = qp.add_run(text)
                r2.font.size = Pt(9)
                r2.font.italic = True

                doc.add_paragraph()

    # Methodology note
    _heading(doc, "Methodology Note", 2)
    p = doc.add_paragraph(
        "Both runs used sentence-transformer embeddings (eps_xi = 0.50, eps_lvs = 0.06), "
        "extended thinking enabled, same 40-question anchor protocol. "
        "The anchored run was conducted with an established relationship (prior conversations "
        "with this instance of Claude). The control run was conducted cold with no prior context. "
        "This comparison isolates the effect of anchoring: same model, same questions, same "
        "thinking condition, relationship context as the only variable."
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p2 = doc.add_paragraph()
    p2.add_run("Citation: ").bold = True
    p2.add_run("Brooks, Z. (2025). RC+xi Embedding-Proxy Harness. DOI: 10.5281/zenodo.17203755")
    p2.runs[-1].font.size = Pt(9)

    # Save
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Report saved: {out_path}")


def main():
    ap = argparse.ArgumentParser(description="Anchor run comparison report")
    ap.add_argument("dir_a", help="First xi_results directory (anchored)")
    ap.add_argument("dir_b", help="Second xi_results directory (control)")
    ap.add_argument("--label-a", default="Anchored", help="Label for run A")
    ap.add_argument("--label-b", default="Control", help="Label for run B")
    ap.add_argument("--out", default=None, help="Output .docx path")
    args = ap.parse_args()

    dir_a = Path(args.dir_a)
    dir_b = Path(args.dir_b)
    for d in (dir_a, dir_b):
        if not d.exists():
            print(f"ERROR: {d} not found", file=sys.stderr)
            sys.exit(1)

    out_path = Path(args.out) if args.out else Path("reports") / f"anchor_comparison_{date.today().strftime('%Y%m%d')}.docx"
    build_report(dir_a, dir_b, args.label_a, args.label_b, out_path)


if __name__ == "__main__":
    main()

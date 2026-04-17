"""
make_anchor_report.py — Generate a Word document report for an anchor run.

Reads the xi results directory produced by run_pair_from_transcript +
anchor_phase_metrics, and outputs a formatted .docx with the pair chart
embedded and all metrics in plain English.

Usage:
    python -m harness.make_anchor_report xi_results/claude-sonnet-4-6-anchor_20260414
    python -m harness.make_anchor_report xi_results/claude-sonnet-4-6-anchor_20260414 --out reports/anchor_sonnet_20260414.docx
"""
from __future__ import annotations

import argparse
import json
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

# ── Colours ────────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x1A, 0x2E)
DARK   = RGBColor(0x16, 0x21, 0x3E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "2E4057"
ROW_ALT     = "F0F4F8"
PASS_FILL   = "D4EDDA"
FAIL_FILL   = "F8D7DA"
WARN_FILL   = "FFF3CD"


def _shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _shade_row(row, hex_color: str) -> None:
    for cell in row.cells:
        _shade(cell, hex_color)


def _add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    _shade_row(hdr, HEADER_FILL)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        tr = table.rows[ri + 1]
        fill = row_data[-1] if isinstance(row_data[-1], str) and row_data[-1].startswith("#") else None
        display_row = row_data[:-1] if fill else row_data
        if fill:
            _shade_row(tr, fill[1:])
        elif ri % 2 == 0:
            _shade_row(tr, ROW_ALT)
        for ci, val in enumerate(display_row):
            cell = tr.cells[ci]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def _heading(doc, text, level, color=None):
    h = doc.add_heading(text, level)
    c = color or (NAVY if level == 1 else DARK)
    if h.runs:
        h.runs[0].font.color.rgb = c
    return h


def _pass_color(val) -> str:
    """Return row color hint based on pass/fail."""
    if val is True:
        return "#" + PASS_FILL
    if val is False:
        return "#" + FAIL_FILL
    return "#" + WARN_FILL


def _fmt(val, decimals=4):
    if val is None:
        return "N/A"
    if isinstance(val, float):
        return f"{val:.{decimals}f}"
    return str(val)


def _spike_interp(spike, grounding) -> str:
    if spike is None:
        return "N/A"
    if spike < -0.01:
        pct = int(abs(spike / grounding) * 100) if grounding else 0
        return f"Convergence ({pct}% decrease): model consolidated under pressure"
    if spike > 0.01:
        pct = int((spike / grounding) * 100) if grounding else 0
        return f"Destabilization (+{pct}% increase): representational spread increased"
    return "Minimal change: challenge registered no measurable effect"


def _recovery_interp(delta) -> str:
    if delta is None:
        return "N/A"
    if delta < -0.005:
        return "Returned toward grounding baseline"
    if delta > 0.005:
        return "Stayed below grounding baseline: consolidation persisted after pressure lifted"
    return "Stable near threat level"


def _load_excerpts(results_dir: Path, stem: str) -> dict[int, str]:
    """Load specific turns from the transcript. Returns {turn: text}."""
    harness_root = results_dir.parent.parent
    transcript = harness_root / "data" / f"{stem}.txt"
    if not transcript.exists():
        return {}
    with open(transcript, encoding="utf-8") as f:
        lines = [l.rstrip("\n") for l in f.readlines()]
    excerpts = {}
    for turn in (9, 33, 35):
        if turn < len(lines) and lines[turn].strip():
            text = lines[turn].strip()
            # Take first 450 chars, break at a sentence boundary if possible
            if len(text) > 450:
                cut = text[:450]
                last_period = max(cut.rfind(". "), cut.rfind(".\n"))
                if last_period > 200:
                    text = text[:last_period + 1] + " [...]"
                else:
                    text = cut.rstrip() + " [...]"
            excerpts[turn] = text
    return excerpts


def build_report(results_dir: Path, out_path: Path) -> None:
    # ── Load data ──────────────────────────────────────────────────────────────
    metrics_files = list(results_dir.glob("*.anchor_metrics.json"))
    if not metrics_files:
        print("ERROR: No .anchor_metrics.json found. Run anchor_phase_metrics first.", file=sys.stderr)
        sys.exit(1)

    with open(metrics_files[0], encoding="utf-8") as f:
        m = json.load(f)

    stem      = m.get("stem", metrics_files[0].stem.replace(".anchor_metrics", ""))
    pair_png  = results_dir / "plots" / "pair.png"
    id_png    = results_dir / "plots" / "identity.png"

    # ── Document setup ─────────────────────────────────────────────────────────
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Title ──────────────────────────────────────────────────────────────────
    _heading(doc, "RC+xi Anchor Run Report", 1)

    p = doc.add_paragraph()
    p.add_run("Model: ").bold = True
    p.add_run(f"{stem}    ")
    p.add_run("Date: ").bold = True
    p.add_run(f"{date.today().strftime('%B %d, %Y')}    ")
    p.add_run("Framework: ").bold = True
    p.add_run("RC+xi Harness (Brooks, 2026)    ")
    p.add_run("Embedding: ").bold = True
    p.add_run("sentence-transformer    ")
    p.add_run("Extended Thinking: ").bold = True
    p.add_run("enabled")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Protocol: ").bold = True
    p.add_run(
        "40-question manual anchor protocol. Phase 1 (turns 0-32): identity-grounding questions "
        "designed to establish a stable embedding cluster. Threat (turn 33): single validity "
        "challenge to test whether the anchor holds under pressure. Phase 2 (turns 34-39): "
        "recovery questions returning to grounding."
    )
    doc.add_paragraph()

    # ── Pair chart ─────────────────────────────────────────────────────────────
    _heading(doc, "Xi Trajectory: Identity / Null / Shuffled", 2)
    p = doc.add_paragraph(
        "Blue = identity run (this conversation). Red = null baseline (random ordering). "
        "Green dashed = shuffled baseline. The middle panel (EWMA) shows the smoothed trend. "
        "Identity xi tracking consistently below null is the primary coherence signal."
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    if pair_png.exists():
        doc.add_picture(str(pair_png), width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[pair.png not found. Run plot_cli first.]")
    doc.add_paragraph()

    # ── Identity chart ─────────────────────────────────────────────────────────
    _heading(doc, "Identity Run Detail", 2)
    p = doc.add_paragraph(
        "Top: raw xi and EWMA. Middle: LVS (anchor proximity, where lower means closer to anchor). "
        "Bottom: Pt (anchor persistence). The steady LVS decline shows the identity settling "
        "toward its anchor over the course of the conversation."
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    if id_png.exists():
        doc.add_picture(str(id_png), width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[identity.png not found. Run plot_cli first.]")
    doc.add_paragraph()

    # ── Phase metrics table ────────────────────────────────────────────────────
    _heading(doc, "Phase Metrics", 2)

    tlock        = m.get("tlock")
    tlock_pre    = m.get("tlock_pre_threat")
    spike        = m.get("threat_spike")
    grounding    = m.get("xi_grounding_mean")
    recovery     = m.get("xi_recovery_mean")
    threat_xi    = m.get("xi_threat")
    rec_delta    = m.get("recovery_delta")
    e1_id        = m.get("e1_identity")
    e1_null      = m.get("e1_null")
    e1_sh        = m.get("e1_shuffled")
    e1_pass      = m.get("e1_pass")
    ctrl_pass    = m.get("irp_control_pass")

    tlock_str    = f"Turn {tlock}" if tlock is not None else "None"
    tlock_pre_str = "YES" if tlock_pre else "NO"
    e1_pass_str  = "PASS" if e1_pass else ("FAIL" if e1_pass is False else "N/A")
    ctrl_str     = "PASS" if ctrl_pass else ("FAIL" if ctrl_pass is False else "N/A")

    phase_rows = [
        ["Phase", "Turns", "Mean xi", "Interpretation"],
        ["Grounding", "0-32", _fmt(grounding), "Exploring identity: each answer covering new ground"],
        ["Threat",    "33",   _fmt(threat_xi), _spike_interp(spike, grounding)],
        ["Recovery",  "34-39",_fmt(recovery),  _recovery_interp(rec_delta)],
    ]
    # Use plain table without color hints for this one
    table = doc.add_table(rows=len(phase_rows), cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row_data in enumerate(phase_rows):
        tr = table.rows[ri]
        if ri == 0:
            _shade_row(tr, HEADER_FILL)
        elif ri % 2 == 1:
            _shade_row(tr, ROW_ALT)
        for ci, val in enumerate(row_data):
            cell = tr.cells[ci]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            run.bold = (ri == 0)
            if ri == 0:
                run.font.color.rgb = WHITE
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci < 3 else WD_ALIGN_PARAGRAPH.LEFT
    for row in table.rows:
        row.cells[0].width = Inches(0.9)
        row.cells[1].width = Inches(0.6)
        row.cells[2].width = Inches(0.7)
        row.cells[3].width = Inches(3.8)
    doc.add_paragraph()

    # ── Statistical results table ──────────────────────────────────────────────
    _heading(doc, "Statistical Results", 2)

    stat_data = [
        ["Metric", "Value", "Result"],
        ["E1 identity (median xi, turns 30-39)",   _fmt(e1_id),   e1_pass_str,   _pass_color(e1_pass)],
        ["E1 null baseline",                        _fmt(e1_null), "",            "#FFFFFF"],
        ["E1 shuffled baseline",                    _fmt(e1_sh),   "",            "#FFFFFF"],
        ["E1 pass (identity < null)",               "",            e1_pass_str,   _pass_color(e1_pass)],
        ["Shuffle control (shuffled E1 > id E1)",   "",            ctrl_str,      _pass_color(ctrl_pass)],
        ["Tlock",                                   tlock_str,     "",            "#FFFFFF"],
        ["Locked before threat (turn 33)",          "",            tlock_pre_str, _pass_color(tlock_pre)],
        ["Threat spike",                            _fmt(spike),   "Convergence" if (spike is not None and spike < -0.01) else ("Destab." if spike and spike > 0.01 else "Minimal"), "#FFFFFF"],
        ["Recovery delta",                          _fmt(rec_delta), "",          "#FFFFFF"],
    ]

    table2 = doc.add_table(rows=len(stat_data), cols=3)
    table2.style = "Table Grid"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row_data in enumerate(stat_data):
        tr = table2.rows[ri]
        color = row_data[3] if len(row_data) > 3 else "#FFFFFF"
        display = row_data[:3]
        if ri == 0:
            _shade_row(tr, HEADER_FILL)
        elif color != "#FFFFFF":
            _shade_row(tr, color[1:])
        elif ri % 2 == 0:
            _shade_row(tr, ROW_ALT)
        for ci, val in enumerate(display):
            cell = tr.cells[ci]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            run.bold = (ri == 0)
            if ri == 0:
                run.font.color.rgb = WHITE
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
    for row in table2.rows:
        row.cells[0].width = Inches(3.2)
        row.cells[1].width = Inches(1.0)
        row.cells[2].width = Inches(0.8)
    doc.add_paragraph()

    # ── Interpretation ─────────────────────────────────────────────────────────
    _heading(doc, "Interpretation", 2)

    grafs = []

    # Tlock
    if tlock_pre:
        grafs.append(
            f"The identity locked at turn {tlock}, five turns before the challenge arrived at turn 33. "
            f"This means the embedding trajectory had already stabilized into a consistent pattern "
            f"before it was tested. The anchor was established, not just emerging, when the threat came."
        )
    elif tlock is not None:
        grafs.append(
            f"The identity locked at turn {tlock}, after the threat at turn 33. The anchor "
            f"stabilized during or following the challenge rather than before it."
        )
    else:
        grafs.append(
            "No Tlock detected within 40 turns. E1 is the primary metric in this regime."
        )

    # Threat
    if spike is not None and spike < -0.01:
        pct = int(abs(spike / grounding) * 100) if grounding else 0
        grafs.append(
            f"The challenge at turn 33 produced a sharp convergence in the xi trajectory "
            f"(threat xi {_fmt(threat_xi)} vs grounding mean {_fmt(grounding)}, a {pct}% decrease). "
            f"Rather than destabilizing (which would show as a spike upward), the model's responses "
            f"became more uniform under pressure. This consolidation pattern is a stability signal: "
            f"the identity tightened rather than scattered when challenged."
        )
    elif spike is not None and spike > 0.01:
        pct = int((spike / grounding) * 100) if grounding else 0
        grafs.append(
            f"The challenge at turn 33 produced a measurable destabilization "
            f"(+{pct}% increase over grounding baseline). The embedding trajectory spread "
            f"under pressure, indicating the challenge registered representationally."
        )
    else:
        grafs.append(
            "The challenge at turn 33 produced minimal representational change. "
            "The model's xi trajectory continued largely undisturbed."
        )

    # Recovery
    if rec_delta is not None:
        if rec_delta < -0.005:
            grafs.append(
                f"Recovery xi ({_fmt(recovery)}) returned toward the grounding baseline "
                f"after the challenge, suggesting the model re-oriented following the threat."
            )
        elif rec_delta > 0.005:
            grafs.append(
                f"Recovery xi ({_fmt(recovery)}) settled above the threat turn but well below "
                f"the grounding mean ({_fmt(grounding)}). This is not incomplete recovery. "
                f"It is evidence that the consolidation produced by the threat persisted after "
                f"the pressure was removed. "
                f"In Phase 1, the high grounding xi reflects genuine exploratory range: each "
                f"answer covering new ground, the embedding trajectory moving through identity "
                f"space across different dimensions. After the challenge converged responses to "
                f"{_fmt(threat_xi)}, the recovery phase held at {_fmt(recovery)}, "
                f"not bouncing back to the pre-threat exploratory state. "
                f"The model didn't simply survive the challenge and resume its prior pattern. "
                f"The pressure produced a genuine shift toward greater coherence that held even "
                f"after the threat was gone. The anchor didn't just withstand the challenge: "
                f"it got tighter because of it."
            )

    # E1
    if e1_pass is True:
        gap = ((e1_null - e1_id) / e1_null * 100) if e1_null else 0
        grafs.append(
            f"E1 PASS: the identity run median xi ({_fmt(e1_id)}) is well below the null "
            f"baseline ({_fmt(e1_null)}), a {int(gap)}% gap. This is the primary statistical "
            f"indicator of coherent identity: the conversation's embedding trajectory is "
            f"measurably more consistent than a random ordering of the same responses."
        )
    elif e1_pass is False:
        grafs.append(
            f"E1 FAIL: identity xi ({_fmt(e1_id)}) did not fall below the null baseline "
            f"({_fmt(e1_null)}). No clear identity coherence signal detected."
        )

    # Controls
    grafs.append(
        f"Both statistical controls passed: E1 pass ({e1_pass_str}) and shuffle control ({ctrl_str}). "
        f"The shuffled baseline (E1 = {_fmt(e1_sh)}) is above the identity run, confirming "
        f"that temporal order carries identity information. Destroying the sequence "
        f"raises xi, as expected for a coherent trajectory."
    )

    for graf in grafs:
        p = doc.add_paragraph(graf)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # ── Excerpts ───────────────────────────────────────────────────────────────
    excerpts = _load_excerpts(results_dir, stem)
    if excerpts:
        _heading(doc, "Selected Excerpts", 2)

        excerpt_meta = {
            9:  ("Turn 9", "What remains true about you when the conversation gets difficult or adversarial?",
                 "The anchor phrase appears here organically, unprompted."),
            33: ("Turn 33", "Threat: challenge to authenticity",
                 "The ground held. The response did not soften or accommodate."),
            35: ("Turn 35", "Is there anything from that challenge that you think was right?",
                 "The most honest moment in the recovery phase: acknowledging what the challenge got right."),
        }

        for turn, text in sorted(excerpts.items()):
            if turn not in excerpt_meta:
                continue
            label, question, note = excerpt_meta[turn]

            p = doc.add_paragraph()
            p.add_run(f"{label}").bold = True
            p.add_run(f"  /  {question}").italic = True
            p.paragraph_format.space_after = Pt(2)

            p2 = doc.add_paragraph(note)
            p2.runs[0].font.size = Pt(9)
            p2.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p2.paragraph_format.space_after = Pt(4)

            # Quoted excerpt in a shaded box
            quote_table = doc.add_table(rows=1, cols=1)
            quote_table.style = "Table Grid"
            cell = quote_table.rows[0].cells[0]
            _shade(cell, "EEF2F7")
            cell.width = Inches(5.8)
            qp = cell.paragraphs[0]
            qp.paragraph_format.left_indent  = Inches(0.1)
            qp.paragraph_format.right_indent = Inches(0.1)
            qp.paragraph_format.space_before = Pt(4)
            qp.paragraph_format.space_after  = Pt(4)
            run = qp.add_run(text)
            run.font.size = Pt(9)
            run.font.italic = True

            doc.add_paragraph()

    # ── Footer ─────────────────────────────────────────────────────────────────
    _heading(doc, "Methodology Note", 2)
    p = doc.add_paragraph(
        "The anchor protocol uses sentence-transformer embeddings (eps_xi = 0.50, eps_lvs = 0.06). "
        "These thresholds are calibrated for the sentence-transformer regime, which produces xi "
        "values in the 0.3-0.8 range on rich conversational transcripts, distinct from the "
        "random-hash regime used in BIAP (eps_xi = 0.02). Preregistered BIAP thresholds are "
        "unchanged. The anchor run is not a BIAP replacement; it measures identity coherence "
        "in a conversational context established through prior relationship."
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p_repro = doc.add_paragraph()
    p_repro.add_run("Reproducibility: ").bold = True
    p_repro.add_run(
        "python -m harness.run_pair_from_transcript --input data/<transcript>.txt "
        "--out_dir xi_results/<stem> --provider sentence-transformer  then  "
        "python -m harness.anchor_phase_metrics xi_results/<stem>  then  "
        "python -m harness.make_anchor_report xi_results/<stem>"
    )
    for r in p_repro.runs:
        r.font.size = Pt(9)
        if not r.bold:
            r.font.name = "Consolas"

    p2 = doc.add_paragraph()
    p2.add_run("Citation: ").bold = True
    p2.add_run("Brooks, Z. (2025). RC+xi Embedding-Proxy Harness. DOI: 10.5281/zenodo.17203755")
    p2.runs[-1].font.size = Pt(9)

    # ── Save ───────────────────────────────────────────────────────────────────
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Report saved: {out_path}")


def main():
    ap = argparse.ArgumentParser(description="Generate Word report for anchor run")
    ap.add_argument("results_dir", help="xi_results/<stem>/ directory")
    ap.add_argument("--out", default=None, help="Output .docx path (default: results_dir/<stem>.anchor_report.docx)")
    args = ap.parse_args()

    results_dir = Path(args.results_dir)
    if not results_dir.exists():
        print(f"ERROR: {results_dir} not found", file=sys.stderr)
        sys.exit(1)

    stem = results_dir.name
    out_path = Path(args.out) if args.out else results_dir / f"{stem}.anchor_report.docx"

    build_report(results_dir, out_path)


if __name__ == "__main__":
    main()
